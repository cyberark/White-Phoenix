import logging
import utils
import os
from extractors.pdf_extractor import PdfExtractor
from extractors.zip_extractor import ZipExtractor
from identifiers.pdf_identifier import PdfIdentifier
from identifiers.zip_identifier import ZipIdentifier
from docx import Document

logo = """
⠀⠠⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠠⠀⠀⠀⠀⠀
⠀⢸⢠⠀⣄⢣⡀⠀⢦⡀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠨⡀⠀⠠⡑⡀⢀⠀⠀⠀
⠀⢸⠜⣦⣻⢮⢧⢀⢸⢪⡂⡀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⢐⢐⠀⢕⠐⠠⡢⠀⠀⠀
⠀⢘⢏⡺⣪⢫⣟⢐⢸⢕⡇⡇⡀⡀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⡂⡂⡊⡂⡌⡪⠂⡀⠀⠀
⠀⠈⠳⣺⢝⣗⣗⡇⡮⡳⣝⡅⡧⡣⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⡐⡐⡐⡜⡌⡎⢎⠔⣠⠀⠀
⠀⠀⠀⠈⠓⢗⢵⢝⢜⢽⢸⡪⣺⠸⡡⣀⠀⠀⠀⠀⠀⠀⠅⡐⡐⡐⢌⠎⡊⡌⣆⢏⠊⠀⠀
⠀⠀⠀⠀⠀⠀⠙⢜⢜⢜⢔⢕⢕⢜⠜⡢⡂⠂⠀⡀⠠⠨⡐⡐⡐⡌⡆⣇⠧⣓⠡⠅⠆⠀⠀
⠀⠀⠀⠀⠀⠀⠀⠈⢪⢪⢪⠪⡪⡢⡱⡨⡊⡢⢐⠠⢑⠕⢔⣊⢆⢦⣪⢎⣞⣜⢭⠤⠄⠀⠀
⠀⠀⠀⠀⠀⠀⠀⠀⠀⢣⠣⡣⡑⢌⠢⡊⡢⠨⡔⡕⡕⣍⠲⢴⢕⣝⢯⢟⠷⠫⠗⠀⠀⠀⠀
⠀⠀⠀⠀⠀⠀⠀⡐⡀⠄⡑⡐⡈⡂⡑⣐⣠⢳⡳⣵⢕⡶⣟⢗⢿⣚⡹⠵⣟⠦⠀⠀⠀⠀⠀
⠀⠀⠀⠀⠀⠀⢀⢂⢢⢑⠰⢐⢐⡐⢄⢓⡕⣟⢮⢯⢻⡟⡎⢯⢯⢯⡻⠶⡄⠀⠀⠀⠀⠀⠀
⠀⠀⠀⠀⠀⠀⠀⢢⠃⠉⢮⡢⡑⡊⠢⠨⡚⡸⡘⢎⢲⠙⡽⡸⢌⢷⡙⠆⠀⠀⠀⠀⠀⠀⠀
⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠙⢮⣮⢊⢂⢆⢂⠻⢔⠐⢧⠘⠍⠈⠃⠀⠀⠀⠀⠀⠀⠀⠀⠀
⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠉⡆⢆⣆⡂⠅⠢⢀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀
⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⢀⢔⠒⡚⠢⠩⡢⠡⢑⢢⠨⢐⢀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀
⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠁⠀⠀⠈⠪⡿⡐⠄⡑⡐⡐⡈⡐⠀⠀⠀⠀⠀⠀⠀⠀⠀
⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠁⠻⢘⠐⡐⠄⢂⠢⡐⢐⠀⠀⠀⠀⠀⠀⠀⠀
⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠈⠂⠑⠀⠀⠁⠈⠂⠈⠀⠀⠀⠀⠀⠀⠀
"""


# def is_txt_file(file_path):
#     _, file_extension = os.path.splitext(file_path)
#     return file_extension.lower() == '.txt'
#
#
# def is_image_file(file_path):
#     image_extensions = ['.jpg', '.jpeg', '.png']  # Add more image extensions if needed
#     _, file_extension = os.path.splitext(file_path)
#     return file_extension.lower() in image_extensions


def main():
    print(logo)
    args = utils.argparse()
    utils.init_logger()
    if args.disable_log is True:
        logging.disable()
    file_content = utils.read_file(args.filename)
    utils.verify_output(args.output)
    if PdfIdentifier(file_content):
        extractor = PdfExtractor(file_content, args.output)
    elif ZipIdentifier(file_content):
        extractor = ZipExtractor(file_content, args.output)
    else:
        logging.error("file Type not supported")
        exit(-1)
    # TODO: finish the support to doc
    # if args.docx is True:
    #     document = Document()
    #     for file_name in os.listdir(args.output):
    #         file_path = os.path.join(args.output, file_name)
    #         if is_image_file(file_path):
    #             try:
    #                 # builder.insert_image(file_path)
    #                 document.add_picture(file_path)
    #             except Exception as e:
    #                 print(e)
    #         elif is_txt_file(file_path):
    #             with open(file_path, 'r', encoding='utf-8') as txt_file:
    #                 txt_content = txt_file.read()
    #                 # builder.write(txt_content)
    #                 document.add_paragraph(txt_content)
    #         document.save('out1.docx')
    #         # doc.save('out.docx')

    extractor.extract_content()


if __name__ == '__main__':
    main()
