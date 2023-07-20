import zlib
import logging
import os
from argparse import ArgumentParser
from io import BytesIO
from PIL import Image
import numpy as np
from docx import Document

logo = """
⠀⠠⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠠⠀⠀⠀⠀⠀
⠀⢸⢠⠀⣄⢣⡀⠀⢦⡀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠨⡀⠀⠠⡑⡀⢀⠀⠀⠀
⠀⢸⠜⣦⣻⢮⢧⢀⢸⢪⡂⡀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⢐⢐⠀⢕⠐⠠⡢⠀⠀⠀
⠀⢘⢏⡺⣪⢫⣟⢐⢸⢕⡇⡇⡀⡀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⡂⡂⡊⡂⡌⡪⠂⡀⠀⠀
⠀⠈⠳⣺⢝⣗⣗⡇⡮⡳⣝⡅⡧⡣⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⡐⡐⡐⡜⡌⡎⢎⠔⣠⠀⠀
⠀⠀⠀⠈⠓⢗⢵⢝⢜⢽⢸⡪⣺⠸⡡⣀⠀⠀⠀⠀⠀⠀⠅⡐⡐⡐⢌⠎⡊⡌⣆⢏⠊⠀⠀
⠀⠀⠀⠀⠀⠀⠙⢜⢜⢜⢔⢕⢕⢜⠜⡢⡂⠂⠀⡀⠠⠨⡐⡐⡐⡌⡆⣇⠧⣓⠡⠅⠆⠀⠀
⠀⠀⠀⠀⠀⠀⠀⠈⢪⢪⢪⠪⡪⡢⡱⡨⡊⡢⢐⠠⢑⠕⢔⣊⢆⢦⣪⢎⣞⣜⢭⠤⠄⠀⠀
⠀⠀⠀⠀⠀⠀⠀⠀⠀⢣⠣⡣⡑⢌⠢⡊⡢⠨⡔⡕⡕⣍⠲⢴⢕⣝⢯⢟⠷⠫⠗⠀⠀⠀⠀
⠀⠀⠀⠀⠀⠀⠀⡐⡀⠄⡑⡐⡈⡂⡑⣐⣠⢳⡳⣵⢕⡶⣟⢗⢿⣚⡹⠵⣟⠦⠀⠀⠀⠀⠀
⠀⠀⠀⠀⠀⠀⢀⢂⢢⢑⠰⢐⢐⡐⢄⢓⡕⣟⢮⢯⢻⡟⡎⢯⢯⢯⡻⠶⡄⠀⠀⠀⠀⠀⠀
⠀⠀⠀⠀⠀⠀⠀⢢⠃⠉⢮⡢⡑⡊⠢⠨⡚⡸⡘⢎⢲⠙⡽⡸⢌⢷⡙⠆⠀⠀⠀⠀⠀⠀⠀
⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠙⢮⣮⢊⢂⢆⢂⠻⢔⠐⢧⠘⠍⠈⠃⠀⠀⠀⠀⠀⠀⠀⠀⠀
⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠉⡆⢆⣆⡂⠅⠢⢀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀
⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⢀⢔⠒⡚⠢⠩⡢⠡⢑⢢⠨⢐⢀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀
⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠁⠀⠀⠈⠪⡿⡐⠄⡑⡐⡐⡈⡐⠀⠀⠀⠀⠀⠀⠀⠀⠀
⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠁⠻⢘⠐⡐⠄⢂⠢⡐⢐⠀⠀⠀⠀⠀⠀⠀⠀
⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠀⠈⠂⠑⠀⠀⠁⠈⠂⠈⠀⠀⠀⠀⠀⠀⠀
"""

document = Document()


def argparse():
    """
    parse the argument to find path of file to extract info from
    :return: arguments
    """
    parser = ArgumentParser(description="Recover text and images from partially encrypted files")
    parser.add_argument("-f", "--file", required=True, dest="filename", metavar="FILE",
                        help="Path to encrypted file")
    parser.add_argument("-o", "--output", required=True, dest="output", metavar="FOLDER",
                        help="Path to folder to save extracted content")
    parser.add_argument("-s", "--separated-files", required=False, dest="separated_files", action='store_true',
                        help="Extract to separate files")
    parser.add_argument("-dl", "--disable-log", required=False, dest="disable_log", action='store_true',
                        help="Disable the log")
    return parser.parse_args()


def verify_output(output_path):
    """
    verify the output folder exists, if not try to create the path
    :param output_path: the output path given as a parameter
    :return:
    """
    if not os.path.exists(output_path):
        try:
            os.mkdir(output_path)
        except OSError as e:
            logging.error(f"Error creating output folder:{e}")
            exit(-1)


def init_logger(disable_logger):
    """
    creates a logger
    :return:
    """
    if disable_logger is True:
        logging.disable()
    else:
        logging.basicConfig(level=logging.INFO, format='%(asctime)s %(levelname)s: %(message)s')
        logging.info(logo)
        logging.info("Copyright © 2023 CyberArk Software Ltd. All rights reserved.\n\n")


def find_object_number(obj_start, content):
    """
    find the number of the object found
    :param obj_start: address of the obj at the start of the object
    :param content: the content of the whole pdf file
    :return: the number of the object found
    """
    object_declaration_row_addr = max(content[: obj_start].rfind(b"\n"), content[: obj_start].rfind(b"\r"))
    object_declaration_row = content[object_declaration_row_addr + 1: obj_start]
    obj_num = int(object_declaration_row.split(b" ")[0])
    logging.info(f"Found object number {obj_num:5} at offset: {hex(object_declaration_row_addr)}")
    return obj_num


def flate_decode(compressed_stream, obj_num):
    """
    decompress extracted stream
    :param compressed_stream: compressed stream, should only contain the actual stream and not the full object
    :param obj_num: the object number of the compressed object
    :return: the decompressed stream or None if fail
    """
    try:
        logging.info(f"Found DEFLATE at object {obj_num}")
        content = zlib.decompress(compressed_stream)
        return content
    except Exception as e:
        logging.error(f"Error decompressing:{e}")
        return None


def read_file(filename):
    """
    read encrypted pdf file
    :param filename: path to encrypted pdf file
    :return: file content
    """
    try:
        with open(filename, "rb") as f:
            content = f.read()
    except FileNotFoundError as e:
        logging.error(f"Error opening file: {e}")
        exit(-1)
    return content


def write_raw_file(image_data, obj_num, output, extension=None):
    """
    save binary as is in file
    :param image_data: the binary of the image
    :param obj_num: the number of the object
    :param output: the output path of the file
    :param extension: the extension of the file
    """
    if extension is None:
        extension = '.jpg'
    image = open(os.path.join(output, str(obj_num)) + extension, "wb")
    image.write(image_data)
    image.close()


def write_to_file(obj_num, file_content, output_path, file_type, separated_files, filter_array=None, mode=None, cmap_len=None):
    """
    write extracted content to file
    :param obj_num: the object from which the content was extracted
    :param file_content: the content of the file to write
    :param output_path: the path to the output folder where the extracted content is to be written
    :param file_type: the type of file image or text
    :param separated_files: boolean for saving in separated files or in docx file
    :param filter_array: the filters of an image
    :param mode: the mode of the image
    :param cmap_len: if text was decoded with cmap, this is the length of the bytes in the cmaps
    :return:
    """
    if file_type == "text":
        if separated_files:
            with open(os.path.join(output_path, str(obj_num) + '.txt'), 'w+') as txt_file:
                txt_file.write(file_content.decode())
        else:
            document.add_paragraph(file_content.decode())
    else:
        temp_file_path = os.path.join('.', 'temp')
        if '/DCTDecode' in filter_array:
            save_jpeg_image(file_content, mode, obj_num, temp_file_path)
        elif '/JPXDecode' in filter_array:
            write_raw_file(file_content, obj_num, temp_file_path, '.jp2')
        else:
            write_raw_file(file_content, obj_num, temp_file_path)
        for file_name in os.listdir(temp_file_path):
            file_path = os.path.join(temp_file_path, file_name)
            if separated_files is not True:
                try:
                    document.add_picture(file_path)
                except Exception as e:
                    logging.error(f'{e} in object number {str(obj_num)}')
                os.remove(file_path)
            else:
                os.replace(file_path, os.path.join(output_path, file_name))
    log = f"Extracted {file_type} content from object {obj_num}" if (cmap_len is None) else \
        f"Extracted {file_type} content from object {obj_num} with cmap from {cmap_len}"
    logging.info(log)


def save_jpeg_image(image_content, mode, obj_num, output):
    """
    get jpeg image object
    :param image_content: the byte array of the image
    :param mode: mode of the image
    :param obj_num: the number of the object
    :param output: the output path of the file
    """
    jpg_data = BytesIO(image_content)
    try:
        image = Image.open(jpg_data)
    except:
        return
    if mode == "CMYK":
        im_data = np.frombuffer(image.tobytes(), dtype='B')
        inv_data = np.full(im_data.shape, 255, dtype='B')
        inv_data -= im_data
        image = Image.frombytes(image.mode, image.size, inv_data.tobytes())
    image.save(output + "//" + str(obj_num) + ".jpg")


def save_doc_file(output):
    document.save(output)
